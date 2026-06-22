#!/usr/bin/env python3
"""Generate StackFix USER_FLOW.docx with rendered diagram images."""

from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "user-flow-assets"
OUTPUT = DOCS / "StackFix_USER_FLOW.docx"

DIAGRAMS: list[tuple[str, str, str]] = [
    (
        "1. Authentication & Session",
        "auth-flow",
        """flowchart TD
  A[Landing / Login] --> B{Valid credentials?}
  B -->|No| C[Show error]
  C --> A
  B -->|Yes| D[Issue JWT + refresh cookie]
  D --> E{User role?}
  E -->|Super Admin| F[Full dashboard + Team]
  E -->|Admin| G[Manager dashboard]
  E -->|Technician| H[My Tickets dashboard]
  F --> I[Idle timeout 8 min]
  G --> I
  H --> I
  I -->|No activity| J[Auto logout]""",
    ),
    (
        "2. Repair Ticket Lifecycle",
        "ticket-lifecycle",
        """stateDiagram-v2
  [*] --> pending: Ticket created
  pending --> under_repair: Start repair
  under_repair --> completed: Repair finished
  completed --> picked_up: Customer collected
  picked_up --> [*]""",
    ),
    (
        "3. New Repair Ticket Flow",
        "new-ticket",
        """flowchart LR
  A[New Repair form] --> B[Customer details]
  B --> C[Device searchable fields]
  C --> D[Fault description]
  D --> E[Submit]
  E --> F[Upsert customer]
  F --> G[Create Pending ticket]
  G --> H[Set Created By]
  H --> I[Ticket detail]""",
    ),
    (
        "4. Invoice & Payment Flow",
        "invoice-flow",
        """flowchart TD
  A[Ticket detail] --> B{Raise invoice?}
  B -->|Admin or Super Admin| C[Create invoice]
  C --> D[Send SMS/WhatsApp]
  D --> E{Payment model}
  E -->|Pay before| F[Pay before Under Repair]
  E -->|Pay on pickup| G[Pay before Picked Up]
  G --> H[Super Admin marks paid]""",
    ),
    (
        "5. Team Management",
        "team-flow",
        """flowchart TD
  A[Team page] --> B[View members]
  B --> C[Add Member modal]
  C --> D[Auto-generate password]
  D --> E[Share with member]
  B --> F[Remove member Super Admin only]""",
    ),
    (
        "6. In-App Messaging",
        "messaging-flow",
        """flowchart TD
  T[Technician] -->|Request edit or delete| M[Messages page]
  M --> A[Admin inbox]
  A -->|Reply| T
  A -->|Resolve| R[Thread closed]""",
    ),
]

PERMISSION_ROWS = [
    ("Create ticket", "Yes", "Yes", "Yes"),
    ("View all tickets", "Yes", "Yes", "Own only"),
    ("Update ticket status", "Yes", "No", "Own only"),
    ("Edit / delete ticket", "Yes", "No", "No"),
    ("Create invoice", "Yes", "Yes", "No"),
    ("View invoices / payments", "Yes", "Yes", "No"),
    ("Mark invoice paid", "Yes", "No", "No"),
    ("Manage team", "Yes", "View only", "No"),
    ("Filter tickets by user", "Yes", "Yes", "No"),
    ("In-app messaging", "Yes", "Yes", "Send requests"),
]

CREDENTIALS = [
    ("Super Admin", "kevin@stackfix.app", "StackFix2026!", "Full CRUD"),
    ("Admin", "admin@stackfix.app", "Admin2026!", "Create + read"),
    ("Technician", "eric@stackfix.app", "Tech2026!", "Own tickets only"),
]


def render_mermaid(name: str, source: str, out_dir: Path) -> Path | None:
    mmd = out_dir / f"{name}.mmd"
    png = out_dir / f"{name}.png"
    mmd.write_text(source.strip() + "\n", encoding="utf-8")
    try:
        subprocess.run(
            ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd), "-o", str(png), "-b", "white"],
            check=True,
            capture_output=True,
            cwd=str(ROOT),
        )
        return png if png.exists() else None
    except (subprocess.CalledProcessError, FileNotFoundError):
        return None


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x1F, 0x12)


def build_doc(diagram_paths: dict[str, Path | None]) -> None:
    doc = Document()
    title = doc.add_heading("StackFix — User Flow Diagrams", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("StackForge AI · Ace Repairs Kigali · Phase 6.3")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph(
        "Visual reference for authentication, roles, repair lifecycle, invoicing, and team flows. "
        "Companion to docs/USER_FLOW.md (markdown + Mermaid source)."
    )

    add_heading(doc, "Demo test accounts", 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Role", "Email", "Password", "Access"]):
        hdr[i].text = label
    for row in CREDENTIALS:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    add_heading(doc, "Role permissions matrix", 2)
    ptable = doc.add_table(rows=1, cols=4)
    ptable.style = "Table Grid"
    ph = ptable.rows[0].cells
    for i, label in enumerate(["Action", "Super Admin", "Admin", "Technician"]):
        ph[i].text = label
    for row in PERMISSION_ROWS:
        cells = ptable.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    add_heading(doc, "Repair status unlock rules", 2)
    doc.add_paragraph(
        "Completed and Picked Up are not selectable from Pending. Follow: "
        "Pending → Under Repair → Completed → Picked Up."
    )
    stable = doc.add_table(rows=1, cols=3)
    stable.style = "Table Grid"
    sh = stable.rows[0].cells
    sh[0].text = "Current status"
    sh[1].text = "Completed available?"
    sh[2].text = "Picked Up available?"
    for row in [
        ("Pending", "No — mark Under Repair first", "No"),
        ("Under Repair", "Yes", "No — mark Completed first"),
        ("Completed", "Yes (current)", "Yes"),
        ("Picked Up", "No (terminal)", "Yes (current)"),
    ]:
        c = stable.add_row().cells
        c[0].text, c[1].text, c[2].text = row

    for title_text, key, _ in DIAGRAMS:
        add_heading(doc, title_text, 2)
        png = diagram_paths.get(key)
        if png and png.exists():
            doc.add_picture(str(png), width=Inches(6.2))
        else:
            doc.add_paragraph("[Diagram: see docs/USER_FLOW.md Mermaid source]")

    add_heading(doc, "Phase roadmap (upcoming)", 2)
    for phase, desc in [
        ("7", "OpenAPI docs, PDF invoices"),
        ("8", "MTN MoMo auto-payment"),
        ("9", "WhatsApp/SMS webhooks"),
        ("10", "Mobile app"),
        ("11", "Playwright E2E"),
        ("12", "Production deploy"),
        ("13", "Beta shop onboarding"),
    ]:
        doc.add_paragraph(f"Phase {phase}: {desc}", style="List Bullet")

    doc.add_paragraph("")
    footer = doc.add_paragraph("hello@stackforgeai.africa · Generated from StackFix monorepo")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


def main() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path | None] = {}
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        for _title, key, source in DIAGRAMS:
            paths[key] = render_mermaid(key, source, tmp_path)
            if paths[key]:
                dest = ASSETS / f"{key}.png"
                dest.write_bytes(paths[key].read_bytes())
                paths[key] = dest
    build_doc(paths)


if __name__ == "__main__":
    main()
